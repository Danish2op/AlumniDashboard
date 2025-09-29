import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
import pandas as pd
import numpy as np
import gspread
import plotly.express as px
import plotly.io as pio
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import seaborn as sns
import tempfile
import os
from datetime import datetime, date, timedelta
from zoneinfo import ZoneInfo
import numpy as np
from io import BytesIO
import gspread
from google.oauth2.service_account import Credentials
import json
from typing import List
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.piecharts import Pie
import base64
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import tempfile
import os
import plotly.io as pio
from PIL import Image as PILImage

# ---------------------------------------------------------
# App configuration
# ---------------------------------------------------------
st.set_page_config(
    page_title="Bootcamp 2025 Analytics Dashboard",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------------------------------------------------------
# Dark Techno Styling
# ---------------------------------------------------------
st.markdown("""
<style>
    /* Main background */
    .stApp {
        background: linear-gradient(135deg, #0c0c0c 0%, #1a1a2e 50%, #16213e 100%);
        color: #ffffff;
    }
    
    /* Headers */
    .main-header {
        font-size: 3.5rem;
        font-weight: 800;
        background: linear-gradient(90deg, #00ffff, #ff00ff, #00ff00);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 0.5rem;
        padding: 1rem;
        text-shadow: 0 0 20px rgba(0, 255, 255, 0.5);
    }
    
    .section-header {
        font-size: 1.8rem;
        font-weight: 700;
        color: #00ffff;
        margin: 2rem 0 1rem 0;
        padding: 0.5rem 1rem;
        border-left: 4px solid #ff00ff;
        background: linear-gradient(90deg, rgba(0, 255, 255, 0.1), transparent);
        border-radius: 8px;
    }
    
    .subsection-header {
        font-size: 1.3rem;
        font-weight: 600;
        color: #00ff00;
        margin: 1.5rem 0 0.5rem 0;
        padding-left: 0.5rem;
        border-left: 3px solid #ff00ff;
    }
    
    /* Metric Cards */
    .metric-card {
        background: rgba(255, 255, 255, 0.05);
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        border: 1px solid rgba(0, 255, 255, 0.3);
        box-shadow: 0 4px 20px rgba(0, 255, 255, 0.2);
        margin: 0.5rem 0;
        transition: all 0.3s ease;
        backdrop-filter: blur(10px);
    }
    
    .metric-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 30px rgba(255, 0, 255, 0.3);
        border-color: #ff00ff;
    }
    
    .metric-value {
        font-size: 2.2rem;
        font-weight: 800;
        margin-bottom: 0.5rem;
        background: linear-gradient(90deg, #00ffff, #ff00ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: #cccccc;
        font-weight: 500;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    .learning-metric {
        border-color: #00ff00;
        box-shadow: 0 4px 20px rgba(0, 255, 0, 0.2);
    }
    
    .learning-metric:hover {
        border-color: #00ff00;
        box-shadow: 0 8px 30px rgba(0, 255, 0, 0.3);
    }
    
    .alumni-metric {
        border-color: #00ffff;
        box-shadow: 0 4px 20px rgba(0, 255, 255, 0.2);
    }
    
    .session-metric {
        border-color: #ff00ff;
        box-shadow: 0 4px 20px rgba(255, 0, 255, 0.2);
    }
    
    .today-metric {
        border-color: #ff6b6b;
        box-shadow: 0 4px 20px rgba(255, 107, 107, 0.2);
    }
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 4px;
        background: rgba(255, 255, 255, 0.05);
        border-radius: 8px;
        padding: 4px;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 45px;
        white-space: pre-wrap;
        background: transparent;
        border-radius: 4px;
        padding: 10px 16px;
        font-weight: 500;
        color: #cccccc;
        border: 1px solid transparent;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #00ffff, #ff00ff);
        color: #000000;
        font-weight: 600;
        border: 1px solid #00ffff;
    }
    
    /* Sidebar */
    .css-1d391kg, .css-1lcbmhc {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
    }
    
    /* Dataframes */
    .dataframe {
        background: rgba(255, 255, 255, 0.05) !important;
        color: white !important;
    }
    
    /* Input widgets */
    .stSlider, .stSelectbox, .stButton>button {
        background: rgba(255, 255, 255, 0.05) !important;
        color: white !important;
        border: 1px solid rgba(0, 255, 255, 0.3) !important;
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, #00ffff, #ff00ff) !important;
        color: black !important;
    }
</style>
""", unsafe_allow_html=True)

# Techno color palette
TECHNO_COLORS = {
    'cyber': ['#00ffff', '#ff00ff', '#00ff00', '#ffff00', '#ff6b6b', '#4ecdc4', '#45b7d1', '#96ceb4'],
    'neon': ['#ff00ff', '#00ffff', '#00ff00', '#ffff00', '#ff0066', '#00ff99', '#0066ff', '#ff6600'],
    'matrix': ['#00ff00', '#00ff99', '#00ff66', '#00cc00', '#009900', '#006600', '#00ff33'],
    'dark': ['#1a1a2e', '#16213e', '#0f3460', '#533483', '#e94560']
}

# Set Plotly dark theme
pio.templates.default = "plotly_dark"

# ---------------------------------------------------------
# Helper functions
# ---------------------------------------------------------
def create_metric_card(value, label, card_type="default"):
    color_class = ""
    if card_type == "learning":
        color_class = "learning-metric"
    elif card_type == "alumni":
        color_class = "alumni-metric"
    elif card_type == "session":
        color_class = "session-metric"
    elif card_type == "today":
        color_class = "today-metric"
    
    card = f"""
    <div class="metric-card {color_class}">
        <div class="metric-value">{value}</div>
        <div class="metric-label">{label}</div>
    </div>
    """
    return card

def parse_session_dates(session_date_str: str, reference_date: date) -> List[date]:
    """
    Parse strings like "29, 30, 01, 02" and handle month rollovers.
    reference_date: provides the base month/year to interpret dates.
    Returns list of date objects.
    """
    if pd.isna(session_date_str) or session_date_str in ['', 'Not Specified']:
        return []
    # Normalize
    text = str(session_date_str).strip()
    parts = [p.strip() for p in text.split(',') if p.strip() != '']
    if not parts:
        return []
    days = []
    for p in parts:
        try:
            # some values might be like '01' or '1'
            days.append(int(float(p)))
        except Exception:
            # try to parse full date if provided like "2024-10-25"
            try:
                dt = datetime.strptime(p.split()[0], "%Y-%m-%d").date()
                days.append(dt.day)
            except Exception:
                continue
    if not days:
        return []
    # Determine month/year rollover:
    # If day sequence decreases (e.g., 29, 30, 01) assume month rolled over once.
    dates = []
    base_month = reference_date.month
    base_year = reference_date.year
    current_month = base_month
    current_year = base_year
    prev_day = -1
    for d in days:
        if prev_day != -1 and d < prev_day:
            # rollover to next month
            if current_month == 12:
                current_month = 1
                current_year += 1
            else:
                current_month += 1
        # sanitize day value (skip invalid days like 0 or >31)
        if 1 <= d <= 31:
            try:
                dates.append(date(current_year, current_month, d))
            except Exception:
                # invalid day for the month (e.g., 31 in a 30-day month), try next month
                try:
                    # attempt adjust by month increment
                    if current_month == 12:
                        alt_month = 1
                        alt_year = current_year + 1
                    else:
                        alt_month = current_month + 1
                        alt_year = current_year
                    dates.append(date(alt_year, alt_month, d))
                except Exception:
                    # give up on this day
                    pass
        prev_day = d
    return dates

def consolidate_session_format_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Detect columns like 'Session Format : One on One' with TRUE/FALSE
    and consolidate into a single 'Session Format' column listing formats.
    """
    # Find columns that indicate session formats
    format_cols = [c for c in df.columns if c.lower().startswith('session format')]
    
    # If the sheet uses the old single column, preserve that too
    if 'Session Format' in df.columns and 'Session Format' not in format_cols:
        # if it's already a combined representation, keep as-is for rows where present
        df['Session Format'] = df['Session Format'].fillna('Not Specified')
    
    if format_cols:
        def row_formats(r):
            fmts = []
            for c in format_cols:
                # Skip the main 'Session Format' column if it exists
                if c == 'Session Format':
                    continue
                
                val = r.get(c, False)
                # Accept various truthy values: TRUE, True, 1, 'TRUE', 'true'
                if isinstance(val, str):
                    v = val.strip().lower()
                    is_true = v in ('true', '1', 'yes', 'y', 't')
                else:
                    is_true = bool(val) and val != 0 and str(val).upper() != 'FALSE'
                
                if is_true:
                    # Extract label after colon if present
                    if ':' in c:
                        label = c.split(':', 1)[1].strip()
                    else:
                        label = c.replace('Session Format', '').replace(':', '').strip()
                        if not label:
                            label = c
                    if label:
                        fmts.append(label)
            return fmts if fmts else ['Not Specified']
        
        # Apply the function to get formats for each row
        df['Session Format'] = df.apply(row_formats, axis=1)
    else:
        # Make sure we normalize the single 'Session Format' column into lists
        if 'Session Format' in df.columns:
            df['Session Format'] = df['Session Format'].fillna('Not Specified').apply(
                lambda x: [s.strip() for s in str(x).split(',') if s.strip()] if x not in [np.nan, '', 'Not Specified'] else ['Not Specified']
            )
        else:
            df['Session Format'] = [['Not Specified']] * len(df)
    
    # Ensure type consistency: each cell should be list of strings
    df['Session Format'] = df['Session Format'].apply(
        lambda x: x if isinstance(x, list) else ([s.strip() for s in str(x).split(',') if s.strip()] or ['Not Specified'])
    )
    
    return df

def calculate_learning_hours(df: pd.DataFrame, today: date = None):
    """
    Calculate learning hours metrics from the dataframe.
    Learning Hours = Number of Attendees × Time Devoted
    Only includes sessions that have already happened (past dates).
    """
    if today is None:
        try:
            today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
        except Exception:
            today = datetime.now().date()
    
    # Initialize learning hours columns
    df['Learning Hours'] = 0
    df['Learning Hours Per Session'] = 0
    
    for idx, row in df.iterrows():
        try:
            # Check if this session has already happened
            session_dates = parse_session_dates(row.get('Session date', ''), reference_date=today)
            has_past_sessions = any(sd < today for sd in session_dates)
            
            # Only calculate learning hours for sessions that have already happened
            if not has_past_sessions:
                df.at[idx, 'Learning Hours'] = 0
                df.at[idx, 'Learning Hours Per Session'] = 0
                continue
            
            # Parse Number of Attendees
            attendees_str = str(row.get('Number of Attendees', '0')).strip()
            if attendees_str in ['', 'nan', 'None']:
                avg_attendees = 0
            elif ',' in attendees_str:
                # Handle multiple values (take average)
                attendees_list = []
                for x in attendees_str.split(','):
                    x = x.strip()
                    if x and x != 'nan':
                        try:
                            attendees_list.append(float(x))
                        except ValueError:
                            continue
                avg_attendees = sum(attendees_list) / len(attendees_list) if attendees_list else 0
            else:
                try:
                    avg_attendees = float(attendees_str)
                except ValueError:
                    avg_attendees = 0
            
            # Parse Time Devoted
            time_str = str(row.get('Time devoted', '0')).strip()
            if time_str in ['', 'nan', 'None']:
                avg_time = 0
            elif ',' in time_str:
                # Handle multiple values (take average)
                time_list = []
                for x in time_str.split(','):
                    x = x.strip()
                    if x and x != 'nan':
                        try:
                            time_list.append(float(x))
                        except ValueError:
                            continue
                avg_time = sum(time_list) / len(time_list) if time_list else 0
            else:
                try:
                    avg_time = float(time_str)
                except ValueError:
                    avg_time = 0
            
            # Calculate learning hours
            learning_hours = avg_attendees * avg_time
            df.at[idx, 'Learning Hours'] = learning_hours
            df.at[idx, 'Learning Hours Per Session'] = learning_hours
            
        except (ValueError, TypeError) as e:
            df.at[idx, 'Learning Hours'] = 0
            df.at[idx, 'Learning Hours Per Session'] = 0
    
    return df

def calculate_learning_metrics(df: pd.DataFrame):
    """
    Calculate comprehensive learning metrics
    """
    # Total learning hours
    total_learning_hours = df['Learning Hours'].sum()
    
    # Learning hours by alumni
    learning_by_alumni = df.groupby('Name')['Learning Hours'].sum().reset_index()
    learning_by_alumni = learning_by_alumni.sort_values('Learning Hours', ascending=False)
    
    # Learning hours by department
    learning_by_department = df.groupby('Department')['Learning Hours'].sum().reset_index()
    learning_by_department = learning_by_department.sort_values('Learning Hours', ascending=False)
    
    # Learning hours by company
    learning_by_company = df.groupby('Company')['Learning Hours'].sum().reset_index()
    learning_by_company = learning_by_company.sort_values('Learning Hours', ascending=False)
    
    # Learning hours by session format
    exploded_df = df.explode('Session Format')
    learning_by_format = exploded_df.groupby('Session Format')['Learning Hours'].sum().reset_index()
    learning_by_format = learning_by_format[learning_by_format['Session Format'] != 'Not Specified']
    learning_by_format = learning_by_format.sort_values('Learning Hours', ascending=False)
    
    # Learning hours by batch
    learning_by_batch = df.groupby('Batch')['Learning Hours'].sum().reset_index()
    learning_by_batch = learning_by_batch.sort_values('Batch', ascending=True)
    
    # Average learning hours per alumni
    avg_learning_per_alumni = total_learning_hours / len(df['Name'].unique()) if len(df['Name'].unique()) > 0 else 0
    
    # Top contributors
    top_contributors = learning_by_alumni.head(10)
    
    # Learning efficiency (learning hours per session)
    avg_learning_per_session = total_learning_hours / len(df) if len(df) > 0 else 0
    
    return {
        'total_learning_hours': total_learning_hours,
        'learning_by_alumni': learning_by_alumni,
        'learning_by_department': learning_by_department,
        'learning_by_company': learning_by_company,
        'learning_by_format': learning_by_format,
        'learning_by_batch': learning_by_batch,
        'avg_learning_per_alumni': avg_learning_per_alumni,
        'avg_learning_per_session': avg_learning_per_session,
        'top_contributors': top_contributors,
        'total_alumni_contributors': len(learning_by_alumni[learning_by_alumni['Learning Hours'] > 0])
    }

def calculate_engagement_metrics(df: pd.DataFrame, today: date):
    """
    Calculate detailed engagement metrics
    """
    # Upcoming sessions (next 7 days)
    upcoming_sessions = []
    for idx, row in df.iterrows():
        session_dates = parse_session_dates(row.get('Session date', ''), reference_date=today)
        for sd in session_dates:
            if today <= sd <= today + timedelta(days=7):
                upcoming_sessions.append({
                    'date': sd,
                    'alumni': row.get('Name', 'Unknown'),
                    'department': row.get('Department', 'Unknown'),
                    'formats': row.get('Session Format', ['Not Specified']),
                    'learning_hours': row.get('Learning Hours Per Session', 0)
                })
    
    # Session distribution by day of week
    all_sessions = []
    for idx, row in df.iterrows():
        session_dates = parse_session_dates(row.get('Session date', ''), reference_date=today)
        for sd in session_dates:
            all_sessions.append({
                'date': sd,
                'day_of_week': sd.strftime('%A'),
                'alumni': row.get('Name', 'Unknown')
            })
    
    sessions_df = pd.DataFrame(all_sessions)
    day_distribution = sessions_df['day_of_week'].value_counts().reindex([
        'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday'
    ], fill_value=0)
    
    # Department engagement
    dept_engagement = df.groupby('Department').agg({
        'Learning Hours': 'sum',
        'Name': 'count'
    }).rename(columns={'Name': 'Alumni Count'}).reset_index()
    
    return {
        'upcoming_sessions': sorted(upcoming_sessions, key=lambda x: x['date']),
        'day_distribution': day_distribution,
        'dept_engagement': dept_engagement,
        'total_upcoming_sessions': len(upcoming_sessions)
    }

# ---------------------------------------------------------
# Enhanced Report Generation with Charts
# ---------------------------------------------------------
def create_chart_images(filtered_df, session_metrics, learning_metrics, engagement_metrics):
    """Create chart images for the reports using matplotlib (no Chrome required)"""
    images = {}
    
    try:
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        from io import BytesIO
        
        # Set matplotlib style
        plt.style.use('seaborn-v0_8')
        colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']
        
        # Session Format Distribution Pie Chart
        exploded_formats = filtered_df.explode('Session Format')['Session Format'].replace('', 'Not Specified')
        exploded_formats = exploded_formats[exploded_formats != 'Not Specified']
        if not exploded_formats.empty:
            format_counts = exploded_formats.value_counts()
            
            # Pie Chart
            fig, ax = plt.subplots(figsize=(8, 6))
            wedges, texts, autotexts = ax.pie(format_counts.values, labels=format_counts.index, 
                                            autopct='%1.1f%%', colors=colors)
            ax.set_title('Session Format Distribution', fontsize=14, fontweight='bold')
            plt.tight_layout()
            
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            images['session_formats_pie'] = buf.read()
            buf.close()
            plt.close()
            
            # Bar Chart
            fig, ax = plt.subplots(figsize=(8, 6))
            bars = ax.barh(format_counts.index, format_counts.values, color=colors[:len(format_counts)])
            ax.set_xlabel('Count')
            ax.set_ylabel('Session Format')
            ax.set_title('Session Format Frequency', fontsize=14, fontweight='bold')
            
            # Add value labels on bars
            for i, bar in enumerate(bars):
                width = bar.get_width()
                ax.text(width + 0.1, bar.get_y() + bar.get_height()/2, 
                        f'{int(width)}', ha='left', va='center')
            
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            images['session_formats_bar'] = buf.read()
            buf.close()
            plt.close()
        
        # Learning Hours by Alumni Bar Chart
        if not learning_metrics['learning_by_alumni'].empty:
            top_alumni = learning_metrics['learning_by_alumni'].head(10)
            
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.barh(top_alumni['Name'], top_alumni['Learning Hours'], 
                        color=plt.cm.viridis(top_alumni['Learning Hours'] / top_alumni['Learning Hours'].max()))
            ax.set_xlabel('Learning Hours')
            ax.set_ylabel('Alumni')
            ax.set_title('Top 10 Alumni by Learning Hours', fontsize=14, fontweight='bold')
            
            # Add value labels
            for i, bar in enumerate(bars):
                width = bar.get_width()
                ax.text(width + 0.1, bar.get_y() + bar.get_height()/2, 
                        f'{width:.1f}', ha='left', va='center')
            
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            images['learning_alumni'] = buf.read()
            buf.close()
            plt.close()
        
        # Department Distribution
        dept_counts = filtered_df['Department'].value_counts()
        if not dept_counts.empty:
            # Department Bar Chart
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.barh(dept_counts.index, dept_counts.values, 
                        color=plt.cm.Blues(dept_counts.values / dept_counts.values.max()))
            ax.set_xlabel('Count')
            ax.set_ylabel('Department')
            ax.set_title('Department Distribution', fontsize=14, fontweight='bold')
            
            # Add value labels
            for i, bar in enumerate(bars):
                width = bar.get_width()
                ax.text(width + 0.1, bar.get_y() + bar.get_height()/2, 
                        f'{int(width)}', ha='left', va='center')
            
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            images['departments_bar'] = buf.read()
            buf.close()
            plt.close()
        
        # Learning by Department
        if not learning_metrics['learning_by_department'].empty:
            dept_learning = learning_metrics['learning_by_department'].head(10)
            
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.barh(dept_learning['Department'], dept_learning['Learning Hours'], 
                        color=plt.cm.Greens(dept_learning['Learning Hours'] / dept_learning['Learning Hours'].max()))
            ax.set_xlabel('Learning Hours')
            ax.set_ylabel('Department')
            ax.set_title('Learning Hours by Department', fontsize=14, fontweight='bold')
            
            # Add value labels
            for i, bar in enumerate(bars):
                width = bar.get_width()
                ax.text(width + 0.1, bar.get_y() + bar.get_height()/2, 
                        f'{width:.1f}', ha='left', va='center')
            
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            images['learning_departments'] = buf.read()
            buf.close()
            plt.close()
        
        # Learning by Batch
        if not learning_metrics['learning_by_batch'].empty:
            batch_data = learning_metrics['learning_by_batch']
            
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(batch_data['Batch'], batch_data['Learning Hours'], 
                    marker='o', linewidth=2, markersize=6, color='#1f77b4')
            ax.set_xlabel('Batch Year')
            ax.set_ylabel('Learning Hours')
            ax.set_title('Learning Hours by Batch Year', fontsize=14, fontweight='bold')
            ax.grid(True, alpha=0.3)
            
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            images['learning_batch'] = buf.read()
            buf.close()
            plt.close()
        
    except Exception as e:
        st.warning(f"⚠️ Could not generate charts for PDF: {e}")
        # Return empty images dict if chart generation fails
        pass
    
    return images
    
    return images

def generate_pdf_report(filtered_df, session_metrics, learning_metrics, engagement_metrics, today):
    """Generate a comprehensive PDF report with all analytics and charts"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch)
    styles = getSampleStyleSheet()
    elements = []
    
    # Generate chart images
    chart_images = create_chart_images(filtered_df, session_metrics, learning_metrics, engagement_metrics)
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,
        textColor=colors.HexColor('#2E86AB')
    )
    elements.append(Paragraph("ALUMNI CONNECT PRO - ANALYTICS REPORT", title_style))
    elements.append(Paragraph(f"Generated on: {today.strftime('%B %d, %Y')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Executive Summary
    elements.append(Paragraph("Executive Summary", styles['Heading2']))
    summary_data = [
        ["Metric", "Value", "Details"],
        ["Total Alumni", len(filtered_df), "All registered alumni"],
        ["Total Learning Hours", f"{learning_metrics['total_learning_hours']:.1f}", "Cumulative impact"],
        ["Alumni Contributors", learning_metrics['total_alumni_contributors'], "Active session hosts"],
        ["Avg Learning per Alumni", f"{learning_metrics['avg_learning_per_alumni']:.1f}", "Average contribution"],
        ["Today's Sessions", session_metrics['today_sessions_count'], "Scheduled for today"],
        ["Upcoming Sessions", engagement_metrics['total_upcoming_sessions'], "Next 7 days"]
    ]
    
    summary_table = Table(summary_data, colWidths=[2*inch, 1.5*inch, 2.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F5F5')),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 20))
    
    # Session Format Analysis with Charts
    elements.append(Paragraph("Session Format Analysis", styles['Heading2']))
    
    # Add session format charts
    if 'session_formats_pie' in chart_images:
        # Save chart to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['session_formats_pie'])
            pie_chart_path = tmpfile.name
        
        # Add pie chart image
        pie_img = Image(pie_chart_path, width=6*inch, height=4*inch)
        elements.append(pie_img)
        elements.append(Spacer(1, 10))
        
        # Add bar chart image
        if 'session_formats_bar' in chart_images:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                tmpfile.write(chart_images['session_formats_bar'])
                bar_chart_path = tmpfile.name
            
            bar_img = Image(bar_chart_path, width=6*inch, height=4*inch)
            elements.append(bar_img)
        
        # Clean up
        os.unlink(pie_chart_path)
        if 'bar_chart_path' in locals():
            os.unlink(bar_chart_path)
    
    elements.append(Spacer(1, 20))
    
    # Top Contributors
    elements.append(Paragraph("Top Learning Contributors", styles['Heading2']))
    top_data = [["Rank", "Alumni", "Learning Hours", "Department"]]
    for i, (_, row) in enumerate(learning_metrics['top_contributors'].head(10).iterrows(), 1):
        alumni_data = filtered_df[filtered_df['Name'] == row['Name']].iloc[0]
        dept = alumni_data.get('Department', 'Not Specified')
        top_data.append([str(i), row['Name'], f"{row['Learning Hours']:.1f}", dept])
    
    top_table = Table(top_data, colWidths=[0.5*inch, 2*inch, 1.5*inch, 2*inch])
    top_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FF6B6B')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    elements.append(top_table)
    
    # Add top contributors chart
    if 'learning_alumni' in chart_images:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['learning_alumni'])
            chart_path = tmpfile.name
        
        elements.append(Spacer(1, 10))
        chart_img = Image(chart_path, width=6*inch, height=4*inch)
        elements.append(chart_img)
        os.unlink(chart_path)
    
    elements.append(Spacer(1, 20))
    
    # Department Analysis
    elements.append(Paragraph("Department Analysis", styles['Heading2']))
    
    # Add department charts
    if 'departments_bar' in chart_images:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['departments_bar'])
            dept_chart_path = tmpfile.name
        
        dept_img = Image(dept_chart_path, width=6*inch, height=4*inch)
        elements.append(dept_img)
        os.unlink(dept_chart_path)
    
    dept_data = [["Department", "Alumni Count", "Learning Hours", "Avg per Alumni"]]
    for _, row in learning_metrics['learning_by_department'].iterrows():
        dept_count = len(filtered_df[filtered_df['Department'] == row['Department']])
        avg_per_alumni = row['Learning Hours'] / dept_count if dept_count > 0 else 0
        dept_data.append([
            row['Department'], 
            str(dept_count), 
            f"{row['Learning Hours']:.1f}", 
            f"{avg_per_alumni:.1f}"
        ])
    
    dept_table = Table(dept_data, colWidths=[2*inch, 1.2*inch, 1.5*inch, 1.3*inch])
    dept_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4ECDC4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    elements.append(dept_table)
    
    # Add learning by department chart
    if 'learning_departments' in chart_images:
        elements.append(Spacer(1, 10))
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['learning_departments'])
            learn_dept_path = tmpfile.name
        
        learn_dept_img = Image(learn_dept_path, width=6*inch, height=4*inch)
        elements.append(learn_dept_img)
        os.unlink(learn_dept_path)
    
    elements.append(Spacer(1, 20))
    
    # Learning Trends
    if 'learning_batch' in chart_images:
        elements.append(Paragraph("Learning Trends by Batch Year", styles['Heading2']))
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['learning_batch'])
            batch_chart_path = tmpfile.name
        
        batch_img = Image(batch_chart_path, width=6*inch, height=4*inch)
        elements.append(batch_img)
        os.unlink(batch_chart_path)
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_word_report(filtered_df, session_metrics, learning_metrics, engagement_metrics, today):
    """Generate a comprehensive Word document report with charts"""
    doc = Document()
    
    # Title
    title = doc.add_heading('ALUMNI CONNECT PRO - ANALYTICS REPORT', 0)
    doc.add_paragraph(f"Generated on: {today.strftime('%B %d, %Y')}")
    doc.add_paragraph()
    
    # Generate chart images
    chart_images = create_chart_images(filtered_df, session_metrics, learning_metrics, engagement_metrics)
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Shading Accent 1'
    
    # Header row
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[0].cells[2].text = "Details"
    
    # Data rows
    metrics_data = [
        ("Total Alumni", str(len(filtered_df)), "All registered alumni"),
        ("Total Learning Hours", f"{learning_metrics['total_learning_hours']:.1f}", "Cumulative impact"),
        ("Alumni Contributors", str(learning_metrics['total_alumni_contributors']), "Active session hosts"),
        ("Avg Learning per Alumni", f"{learning_metrics['avg_learning_per_alumni']:.1f}", "Average contribution"),
        ("Today's Sessions", str(session_metrics['today_sessions_count']), "Scheduled for today"),
        ("Upcoming Sessions", str(engagement_metrics['total_upcoming_sessions']), "Next 7 days")
    ]
    
    for i, (metric, value, details) in enumerate(metrics_data, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value
        table.rows[i].cells[2].text = details
    
    doc.add_paragraph()
    
    # Session Format Analysis with Charts
    doc.add_heading('Session Format Analysis', level=1)
    
    # Add session format charts
    if 'session_formats_pie' in chart_images:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['session_formats_pie'])
            doc.add_picture(tmpfile.name, width=Inches(6))
            os.unlink(tmpfile.name)
        
        doc.add_paragraph()
        
        if 'session_formats_bar' in chart_images:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                tmpfile.write(chart_images['session_formats_bar'])
                doc.add_picture(tmpfile.name, width=Inches(6))
                os.unlink(tmpfile.name)
    
    doc.add_paragraph()
    
    # Top Contributors
    doc.add_heading('Top Learning Contributors', level=1)
    
    # Add top contributors chart
    if 'learning_alumni' in chart_images:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['learning_alumni'])
            doc.add_picture(tmpfile.name, width=Inches(6))
            os.unlink(tmpfile.name)
        
        doc.add_paragraph()
    
    top_table = doc.add_table(rows=len(learning_metrics['top_contributors'].head(10))+1, cols=4)
    top_table.style = 'Light List Accent 2'
    
    top_table.rows[0].cells[0].text = "Rank"
    top_table.rows[0].cells[1].text = "Alumni"
    top_table.rows[0].cells[2].text = "Learning Hours"
    top_table.rows[0].cells[3].text = "Department"
    
    for i, (_, row) in enumerate(learning_metrics['top_contributors'].head(10).iterrows(), 1):
        alumni_data = filtered_df[filtered_df['Name'] == row['Name']].iloc[0]
        dept = alumni_data.get('Department', 'Not Specified')
        top_table.rows[i].cells[0].text = str(i)
        top_table.rows[i].cells[1].text = row['Name']
        top_table.rows[i].cells[2].text = f"{row['Learning Hours']:.1f}"
        top_table.rows[i].cells[3].text = dept
    
    doc.add_paragraph()
    
    # Department Analysis
    doc.add_heading('Department Analysis', level=1)
    
    # Add department charts
    if 'departments_bar' in chart_images:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['departments_bar'])
            doc.add_picture(tmpfile.name, width=Inches(6))
            os.unlink(tmpfile.name)
        
        doc.add_paragraph()
        
        if 'learning_departments' in chart_images:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                tmpfile.write(chart_images['learning_departments'])
                doc.add_picture(tmpfile.name, width=Inches(6))
                os.unlink(tmpfile.name)
            
            doc.add_paragraph()
    
    dept_table = doc.add_table(rows=len(learning_metrics['learning_by_department'])+1, cols=4)
    dept_table.style = 'Light List Accent 3'
    
    dept_table.rows[0].cells[0].text = "Department"
    dept_table.rows[0].cells[1].text = "Alumni Count"
    dept_table.rows[0].cells[2].text = "Learning Hours"
    dept_table.rows[0].cells[3].text = "Avg per Alumni"
    
    for i, (_, row) in enumerate(learning_metrics['learning_by_department'].iterrows(), 1):
        dept_count = len(filtered_df[filtered_df['Department'] == row['Department']])
        avg_per_alumni = row['Learning Hours'] / dept_count if dept_count > 0 else 0
        dept_table.rows[i].cells[0].text = row['Department']
        dept_table.rows[i].cells[1].text = str(dept_count)
        dept_table.rows[i].cells[2].text = f"{row['Learning Hours']:.1f}"
        dept_table.rows[i].cells[3].text = f"{avg_per_alumni:.1f}"
    
    doc.add_paragraph()
    
    # Learning Trends
    if 'learning_batch' in chart_images:
        doc.add_heading('Learning Trends by Batch Year', level=1)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(chart_images['learning_batch'])
            doc.add_picture(tmpfile.name, width=Inches(6))
            os.unlink(tmpfile.name)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# Google Sheets Configuration from secrets.toml
# ---------------------------------------------------------
def get_google_sheets_config():
    """Get Google Sheets configuration from secrets.toml"""
    try:
        # Check if secrets are available
        if 'google_sheets_service_account' not in st.secrets:
            st.error("❌ Google Sheets credentials not found in secrets.toml")
            return None, None
        
        # Build credentials dictionary from secrets
        credentials_dict = {
            "type": st.secrets["google_sheets_service_account"]["type"],
            "project_id": st.secrets["google_sheets_service_account"]["project_id"],
            "private_key_id": st.secrets["google_sheets_service_account"]["private_key_id"],
            "private_key": st.secrets["google_sheets_service_account"]["private_key"],
            "client_email": st.secrets["google_sheets_service_account"]["client_email"],
            "client_id": st.secrets["google_sheets_service_account"]["client_id"],
            "auth_uri": st.secrets["google_sheets_service_account"]["auth_uri"],
            "token_uri": st.secrets["google_sheets_service_account"]["token_uri"],
            "auth_provider_x509_cert_url": st.secrets["google_sheets_service_account"]["auth_provider_x509_cert_url"],
            "client_x509_cert_url": st.secrets["google_sheets_service_account"]["client_x509_cert_url"],
            "universe_domain": st.secrets["google_sheets_service_account"]["universe_domain"]
        }
        
        # Get sheet URL
        sheet_url = st.secrets["google_sheets"]["sheet_url"]
        
        return credentials_dict, sheet_url
        
    except Exception as e:
        st.error(f"❌ Error reading secrets: {e}")
        st.write("Available secrets sections:", list(st.secrets.keys()))
        return None, None

def load_google_sheets(credentials_json: dict, sheet_url: str, worksheet_name: str = "Sheet1"):
    try:
        scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
        creds = Credentials.from_service_account_info(credentials_json, scopes=scope)
        client = gspread.authorize(creds)
        sheet = client.open_by_url(sheet_url)
        worksheet = sheet.worksheet(worksheet_name)
        
        # Get all values as raw data first
        all_values = worksheet.get_all_values()
        
        if not all_values:
            st.error("❌ No data found in the worksheet")
            return None
        
        # Get headers and fix duplicates/empty names
        headers = all_values[0]
        fixed_headers = []
        header_counts = {}
        
        for i, header in enumerate(headers):
            # Clean up the header
            clean_header = str(header).strip()
            if not clean_header or clean_header.lower() in ['false', 'true', '']:
                clean_header = f"Column_{i+1}"
            
            # Handle duplicates
            if clean_header in header_counts:
                header_counts[clean_header] += 1
                clean_header = f"{clean_header}_{header_counts[clean_header]}"
            else:
                header_counts[clean_header] = 0
            
            fixed_headers.append(clean_header)
        
        # Create DataFrame with fixed headers
        data_rows = all_values[1:]  # Skip header row
        df = pd.DataFrame(data_rows, columns=fixed_headers)
        
        # Remove completely empty columns
        df = df.dropna(axis=1, how='all')
        
        # Remove completely empty rows
        df = df.dropna(axis=0, how='all')
        
        return df
        
    except Exception as e:
        st.error(f"❌ Error loading Google Sheets: {e}")
        return None

# ---------------------------------------------------------
# Data cleaning
# ---------------------------------------------------------
def clean_dataframe(df: pd.DataFrame, today: date = None) -> pd.DataFrame:
    if today is None:
        try:
            today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
        except Exception:
            today = datetime.now().date()
    
    # Standardize column names (strip whitespace)
    df.columns = [c.strip() for c in df.columns]
    
    # Fix common typos
    if 'Depatrment' in df.columns and 'Department' not in df.columns:
        df['Department'] = df['Depatrment']
        df.drop(columns=['Depatrment'], inplace=True, errors='ignore')
    
    # Ensure defaults exist
    required_columns = ['Name', 'Batch', 'Department', 'Position', 'Company', 'Format', 'Session date']
    for col in required_columns:
        if col not in df.columns:
            df[col] = 'Not Specified'
    
    # Clean up empty values and replace with 'Not Specified'
    for col in required_columns:
        df[col] = df[col].fillna('Not Specified')
        df[col] = df[col].astype(str).replace('', 'Not Specified')
        df[col] = df[col].replace('nan', 'Not Specified')
    
    # Ensure learning-related columns exist
    if 'Number of Attendees' not in df.columns:
        df['Number of Attendees'] = 0
    
    if 'Time devoted' not in df.columns:
        df['Time devoted'] = 0
    
    # Clean numeric columns
    df['Number of Attendees'] = df['Number of Attendees'].fillna(0)
    df['Time devoted'] = df['Time devoted'].fillna(0)
    
    # Batch numeric - handle empty/invalid values better
    df['Batch'] = df['Batch'].replace('', '2000')  # Default year for empty batches
    df['Batch'] = pd.to_numeric(df['Batch'], errors='coerce').fillna(2000).astype(int)
    
    # Normalize Format column
    if 'Format' not in df.columns:
        df['Format'] = 'Not Specified'
    df['Format'] = df['Format'].fillna('Not Specified').astype(str)
    df['Format'] = df['Format'].replace('', 'Not Specified')
    
    # Consolidate Session Format columns into a list per row
    df = consolidate_session_format_columns(df)
    
    # Normalize Session date
    df['Session date'] = df['Session date'].fillna('').astype(str)
    
    # Calculate learning hours (only for past sessions)
    df = calculate_learning_hours(df, today)
    
    return df

# ---------------------------------------------------------
# Data loader with caching
# ---------------------------------------------------------
@st.cache_data(ttl=600)
def load_data():
    """Load data from Google Sheets using secrets.toml"""
    try:
        # Get today's date
        try:
            today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
        except Exception:
            today = datetime.now().date()
            
        # Get configuration from secrets
        credentials_dict, sheet_url = get_google_sheets_config()
        
        if credentials_dict is None or sheet_url is None:
            return pd.DataFrame()
        
        df = load_google_sheets(credentials_dict, sheet_url)
        if df is not None and not df.empty:
            return clean_dataframe(df, today)
        else:
            st.error("❌ No data found in Google Sheets")
            return pd.DataFrame()
    except Exception as e:
        st.error(f"❌ Error loading data: {e}")
        return pd.DataFrame()

# ---------------------------------------------------------
# Session metrics calculation
# ---------------------------------------------------------
def calculate_session_metrics(df: pd.DataFrame, today: date):
    all_sessions = []
    session_types_today = []
    alumni_hosted_today = []
    for idx, row in df.iterrows():
        session_dates = parse_session_dates(row.get('Session date', ''), reference_date=today)
        session_formats = row.get('Session Format', ['Not Specified'])
        # Ensure we have a list
        if not isinstance(session_formats, list):
            session_formats = [s.strip() for s in str(session_formats).split(',') if s.strip()]
            if not session_formats:
                session_formats = ['Not Specified']
        for sd in session_dates:
            all_sessions.append({
                'date': sd,
                'alumni': row.get('Name', 'Unknown'),
                'department': row.get('Department', 'Unknown'),  # Fixed: use 'department' key
                'formats': [fmt for fmt in session_formats if fmt not in [None, '', 'Not Specified']],
                'learning_hours': row.get('Learning Hours Per Session', 0)
            })
            if sd == today:
                alumni_hosted_today.append(row.get('Name', 'Unknown'))
                session_types_today.extend([fmt for fmt in session_formats if fmt not in [None, '', 'Not Specified']])
    total_alumni_hosted = len(set(alumni_hosted_today))
    unique_session_types_today = len(set(session_types_today))
    today_sessions_count = len([s for s in all_sessions if s['date'] == today])
    remaining_sessions = len([s for s in all_sessions if s['date'] > today])
    return {
        'alumni_hosted_today': total_alumni_hosted,
        'session_types_today': unique_session_types_today,
        'today_sessions_count': today_sessions_count,
        'remaining_sessions': remaining_sessions,
        'all_sessions': all_sessions,
        'today_session_details': [s for s in all_sessions if s['date'] == today]
    }

# ---------------------------------------------------------
# Main Application
# ---------------------------------------------------------
def main():
    # --- CHANGED: New header with a different centered layout ---
    header_html = """
    <div style="text-align: center; margin-bottom: 2rem;">
        <img src="https://i.ibb.co/gh7mjqb/Whats-App-Image-2025-09-25-at-12-33-43-991b6071.jpg" alt="Logo 1" style="height: 80px; margin-bottom: 1rem;">
        <div class="main-header" style="padding: 0; margin-bottom: 0.5rem; font-size: 2.8rem;">Alumni Bootcamp 2025 Analytics Dashboard</div>
        <div style='font-size: 1.2rem; color: #cccccc; margin-bottom: 0.5rem;'>Built For Alumni Relations Office By Team SAIC</div>
        <img src="https://i.ibb.co/V095yskk/saic-logo-png.png" alt="Logo 2" style="height: 40px;">
    </div>
    """
    st.markdown(header_html, unsafe_allow_html=True)

    # Get today's date
    try:
        today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
    except Exception:
        today = datetime.now().date()
    
    # Display today's date prominently
    st.markdown(f"<div style='text-align: center; font-size: 1.1rem; color: #00ffff; margin-bottom: 1rem; background: rgba(0, 255, 255, 0.1); padding: 0.5rem; border-radius: 8px;'>📅 Today: {today.strftime('%A, %B %d, %Y')}</div>", unsafe_allow_html=True)

    # Load data
    with st.spinner('🔄 CONNECTING TO DATA SOURCE...'):
        df = load_data()

    if not df.empty:
        # Sidebar - Filters
        with st.sidebar:
            st.markdown("### 🔍 DATA FILTERS")
            st.sidebar.info(f"📊 {len(df)} ALUMNI RECORDS LOADED")
            
            # Batch slider
            batch_min = int(df['Batch'].min()) if len(df) > 0 else 2000
            batch_max = int(df['Batch'].max()) if len(df) > 0 else datetime.now().year
            batch_range = st.slider("🎓 BATCH RANGE", min_value=batch_min, max_value=batch_max, value=(batch_min, batch_max))

            departments = ['All'] + sorted(df['Department'].fillna('Not Specified').unique().tolist())
            selected_department = st.selectbox("🏫 DEPARTMENT", departments)

            formats = ['All'] + sorted(df['Format'].fillna('Not Specified').unique().tolist())
            selected_format = st.selectbox("💻 FORMAT", formats)

            companies = ['All'] + sorted([c for c in df['Company'].fillna('Not Specified').unique().tolist() if c != 'Not Specified'])
            selected_company = st.selectbox("🏢 COMPANY", companies)

            st.markdown("---")
            st.markdown("### 📊 DATA SOURCE")
            st.success("🔐 SECURELY CONNECTED VIA SECRETS.TOML")
            credentials_dict, sheet_url = get_google_sheets_config()
            if sheet_url:
                st.markdown(f"[🔗 VIEW SOURCE SHEET]({sheet_url})")

        # Apply filters
        filtered_df = df[
            (df['Batch'] >= batch_range[0]) &
            (df['Batch'] <= batch_range[1])
        ].copy()

        if selected_department != 'All':
            filtered_df = filtered_df[filtered_df['Department'] == selected_department]
        if selected_format != 'All':
            filtered_df = filtered_df[filtered_df['Format'] == selected_format]
        if selected_company != 'All':
            filtered_df = filtered_df[filtered_df['Company'] == selected_company]

        # Today's date
        try:
            today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
        except Exception:
            today = datetime.now().date()

        # Compute all metrics
        session_metrics = calculate_session_metrics(filtered_df, today=today)
        learning_metrics = calculate_learning_metrics(filtered_df)
        engagement_metrics = calculate_engagement_metrics(filtered_df, today=today)

        # SECTION: EXECUTIVE OVERVIEW
        st.markdown('<div class="section-header">📊 EXECUTIVE OVERVIEW</div>', unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(create_metric_card(session_metrics['alumni_hosted_today'], "ALUMNI HOSTED TODAY", "alumni"), unsafe_allow_html=True)
        with col2:
            st.markdown(create_metric_card(session_metrics['session_types_today'], "SESSION TYPES TODAY", "session"), unsafe_allow_html=True)
        with col3:
            st.markdown(create_metric_card(session_metrics['today_sessions_count'], "TODAY'S SESSIONS", "today"), unsafe_allow_html=True)
        with col4:
            st.markdown(create_metric_card(f"{learning_metrics['total_learning_hours']:.0f}", "TOTAL LEARNING HOURS", "learning"), unsafe_allow_html=True)

        # SECTION: TODAY'S SESSION DETAILS
        st.markdown('<div class="section-header">📅 TODAY\'S SESSION DETAILS</div>', unsafe_allow_html=True)
        
        if session_metrics['today_session_details']:
            today_sessions_df = pd.DataFrame(session_metrics['today_session_details'])
            today_sessions_df['formats'] = today_sessions_df['formats'].apply(lambda x: ', '.join(x) if isinstance(x, list) else str(x))
            
            st.metric("Sessions Scheduled for Today", len(session_metrics['today_session_details']))
            
            # Display today's sessions in a nice format
            for idx, session in enumerate(today_sessions_df.to_dict('records')):
                # Get department info from the filtered_df
                dept_info = filtered_df[filtered_df['Name'] == session['alumni']]['Department'].iloc[0] if len(filtered_df[filtered_df['Name'] == session['alumni']]) > 0 else "Unknown"
                
                with st.expander(f"Session {idx+1}: {session['alumni']} ({dept_info})"):
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Alumni", session['alumni'])
                    with col2:
                        st.metric("Department", dept_info)
                    with col3:
                        st.metric("Formats", session['formats'])
                    with col4:
                        st.metric("Learning Hours", f"{session['learning_hours']:.1f}")
        else:
            st.info("No sessions scheduled for today")

        # SECTION: INDIVIDUAL ALUMNI ANALYSIS
        st.markdown('<div class="section-header">👤 INDIVIDUAL ALUMNI ANALYSIS</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            # Dropdown to select alumni
            alumni_list = sorted(filtered_df['Name'].unique().tolist())
            selected_alumni = st.selectbox("Select Alumni for Detailed Analysis", alumni_list)
            
            if selected_alumni:
                alumni_data = filtered_df[filtered_df['Name'] == selected_alumni]
                
                # Calculate actual sessions and learning hours by parsing individual data
                total_actual_sessions = 0
                total_actual_learning_hours = 0
                
                for idx, row in alumni_data.iterrows():
                    # Parse attendees and time to count actual sessions
                    attendees_str = str(row.get('Number of Attendees', '0')).strip()
                    time_str = str(row.get('Time devoted', '0')).strip()
                    
                    # Parse attendees list
                    if ',' in attendees_str:
                        attendees_list = []
                        for x in attendees_str.split(','):
                            x = x.strip()
                            if x and x not in ['nan', '', 'None']:
                                try:
                                    attendees_list.append(float(x))
                                except ValueError:
                                    pass
                    else:
                        try:
                            attendees_list = [float(attendees_str)] if attendees_str not in ['', 'nan', 'None'] else []
                        except ValueError:
                            attendees_list = []
                    
                    # Parse time list
                    if ',' in time_str:
                        time_list = []
                        for x in time_str.split(','):
                            x = x.strip()
                            if x and x not in ['nan', '', 'None']:
                                try:
                                    time_list.append(float(x))
                                except ValueError:
                                    pass
                    else:
                        try:
                            time_list = [float(time_str)] if time_str not in ['', 'nan', 'None'] else []
                        except ValueError:
                            time_list = []
                    
                    # Count actual sessions (max of attendees or time entries)
                    session_count = max(len(attendees_list), len(time_list))
                    total_actual_sessions += session_count
                    
                    # Calculate learning hours for each session
                    for i in range(session_count):
                        attendee_count = attendees_list[i] if i < len(attendees_list) else (attendees_list[0] if attendees_list else 0)
                        time_devoted = time_list[i] if i < len(time_list) else (time_list[0] if time_list else 0)
                        total_actual_learning_hours += attendee_count * time_devoted
                
                avg_learning_per_session = total_actual_learning_hours / total_actual_sessions if total_actual_sessions > 0 else 0
                
                # Display alumni metrics
                st.markdown(f"### {selected_alumni}")
                st.metric("Total Sessions", total_actual_sessions)
                st.metric("Total Learning Hours", f"{total_actual_learning_hours:.1f}")
                # st.metric("Avg Hours per Session", f"{avg_learning_per_session:.1f}")
                
                # Alumni details
                if not alumni_data.empty:
                    first_record = alumni_data.iloc[0]
                    st.markdown("**Details:**")
                    st.write(f"**Department:** {first_record['Department']}")
                    st.write(f"**Company:** {first_record['Company']}")
                    st.write(f"**Batch:** {first_record['Batch']}")
        
        with col2:
            if selected_alumni:
                alumni_sessions = filtered_df[filtered_df['Name'] == selected_alumni]
                
                if not alumni_sessions.empty and alumni_sessions['Learning Hours'].sum() > 0:
                    # Create comprehensive session format analysis with corrected calculations
                    format_hours = {}
                    
                    for idx, row in alumni_sessions.iterrows():
                        formats = row['Session Format'] if isinstance(row['Session Format'], list) else [str(row['Session Format'])]
                        
                        # Get individual session data for proper calculation
                        attendees_str = str(row.get('Number of Attendees', '0')).strip()
                        time_str = str(row.get('Time devoted', '0')).strip()
                        
                        # Parse attendees - could be comma separated like "15,20"
                        if ',' in attendees_str:
                            attendees_list = []
                            for x in attendees_str.split(','):
                                x = x.strip()
                                if x and x != 'nan':
                                    try:
                                        attendees_list.append(float(x))
                                    except ValueError:
                                        pass
                        else:
                            try:
                                attendees_list = [float(attendees_str)] if attendees_str not in ['', 'nan', 'None'] else [0]
                            except ValueError:
                                attendees_list = [0]
                        
                        # Parse time - could be comma separated like "0.5,0.5"
                        if ',' in time_str:
                            time_list = []
                            for x in time_str.split(','):
                                x = x.strip()
                                if x and x != 'nan':
                                    try:
                                        time_list.append(float(x))
                                    except ValueError:
                                        pass
                        else:
                            try:
                                time_list = [float(time_str)] if time_str not in ['', 'nan', 'None'] else [0]
                            except ValueError:
                                time_list = [0]
                        
                        # Filter out non-specified formats
                        valid_formats = [fmt for fmt in formats if fmt != 'Not Specified' and fmt.strip()]
                        
                        # If no valid formats but we have learning hours, use "General Session"
                        if not valid_formats and (attendees_list and attendees_list[0] > 0) and (time_list and time_list[0] > 0):
                            valid_formats = ['General Session']
                        
                        if not valid_formats:
                            continue
                        
                        # Calculate learning hours for each format based on actual data
                        # Match formats with their corresponding attendees and time
                        for i, fmt in enumerate(valid_formats):
                            # Use corresponding attendee count and time if available, otherwise use first values
                            attendee_count = attendees_list[i] if i < len(attendees_list) else (attendees_list[0] if attendees_list else 0)
                            time_devoted = time_list[i] if i < len(time_list) else (time_list[0] if time_list else 0)
                            
                            # Calculate actual learning hours for this specific format
                            format_learning_hours = attendee_count * time_devoted
                            
                            format_hours[fmt] = format_hours.get(fmt, 0) + format_learning_hours
                    
                    if format_hours:
                        # Create pie chart showing learning hours by format
                        formats = list(format_hours.keys())
                        hours = list(format_hours.values())
                        
                        fig_pie = px.pie(values=hours, names=formats, hole=0.4,
                                        color_discrete_sequence=TECHNO_COLORS['cyber'], 
                                        title=f"{selected_alumni}'s Learning Hours by Session Format")
                        fig_pie.update_traces(textposition='inside', textinfo='percent+label+value')
                        fig_pie.update_layout(height=500)
                        st.plotly_chart(fig_pie, use_container_width=True)
                        
                        # Add bar chart for alumni engagement (learning hours by alumni)
                        st.markdown("### Alumni Engagement - Learning Hours Comparison")
                        
                        # Calculate actual learning hours for all alumni using the same logic as individual analysis
                        all_alumni_actual_hours = []
                        for alumni_name in filtered_df['Name'].unique():
                            alumni_rows = filtered_df[filtered_df['Name'] == alumni_name]
                            total_alumni_learning_hours = 0
                            
                            for idx, row in alumni_rows.iterrows():
                                # Parse attendees and time to calculate actual learning hours
                                attendees_str = str(row.get('Number of Attendees', '0')).strip()
                                time_str = str(row.get('Time devoted', '0')).strip()
                                
                                # Parse attendees list
                                if ',' in attendees_str:
                                    attendees_list = []
                                    for x in attendees_str.split(','):
                                        x = x.strip()
                                        if x and x not in ['nan', '', 'None']:
                                            try:
                                                attendees_list.append(float(x))
                                            except ValueError:
                                                pass
                                else:
                                    try:
                                        attendees_list = [float(attendees_str)] if attendees_str not in ['', 'nan', 'None'] else []
                                    except ValueError:
                                        attendees_list = []
                                
                                # Parse time list
                                if ',' in time_str:
                                    time_list = []
                                    for x in time_str.split(','):
                                        x = x.strip()
                                        if x and x not in ['nan', '', 'None']:
                                            try:
                                                time_list.append(float(x))
                                            except ValueError:
                                                pass
                                else:
                                    try:
                                        time_list = [float(time_str)] if time_str not in ['', 'nan', 'None'] else []
                                    except ValueError:
                                        time_list = []
                                
                                # Calculate learning hours for each session
                                session_count = max(len(attendees_list), len(time_list))
                                for i in range(session_count):
                                    attendee_count = attendees_list[i] if i < len(attendees_list) else (attendees_list[0] if attendees_list else 0)
                                    time_devoted = time_list[i] if i < len(time_list) else (time_list[0] if time_list else 0)
                                    total_alumni_learning_hours += attendee_count * time_devoted
                            
                            all_alumni_actual_hours.append({
                                'Name': alumni_name,
                                'Learning Hours': total_alumni_learning_hours
                            })
                        
                        # Convert to DataFrame and sort
                        all_alumni_hours = pd.DataFrame(all_alumni_actual_hours)
                        all_alumni_hours = all_alumni_hours[all_alumni_hours['Learning Hours'] > 0]  # Only show alumni with learning hours
                        all_alumni_hours = all_alumni_hours.sort_values('Learning Hours', ascending=False) # Sort descending for better viz

                        # ---- CHANGED: Replaced px.bar with go.Figure to fix hover text ----
                        if not all_alumni_hours.empty:
                            # Separate the selected alumni from the rest
                            selected_alumni_df = all_alumni_hours[all_alumni_hours['Name'] == selected_alumni]
                            other_alumni_df = all_alumni_hours[all_alumni_hours['Name'] != selected_alumni]

                            fig_bar = go.Figure()

                            # Add trace for other alumni with a colorscale
                            fig_bar.add_trace(go.Bar(
                                y=other_alumni_df['Name'],
                                x=other_alumni_df['Learning Hours'],
                                name='Other Alumni',
                                orientation='h',
                                marker=dict(
                                    color=other_alumni_df['Learning Hours'], # Set color data
                                    colorscale='Viridis' # Apply colorscale
                                )
                            ))

                            # Add trace for the selected alumni
                            fig_bar.add_trace(go.Bar(
                                y=selected_alumni_df['Name'],
                                x=selected_alumni_df['Learning Hours'],
                                name=selected_alumni,
                                orientation='h',
                                marker=dict(
                                    color='#ff00ff' # Highlight color
                                )
                            ))

                            # Update layout
                            fig_bar.update_layout(
                                title_text="Learning Hours by Alumni (Selected alumni highlighted)",
                                height=max(400, len(all_alumni_hours) * 25),
                                xaxis_title="Learning Hours",
                                yaxis_title="Alumni",
                                template="plotly_dark",
                                yaxis={'categoryorder':'total descending'},
                                showlegend=False
                            )
                            
                            st.plotly_chart(fig_bar, use_container_width=True)
                        else:
                            st.info("No alumni with recorded learning hours found")
                    else:
                        st.info("No session format data available for this alumni")
                else:
                    st.info("No learning hours recorded for this alumni")

        # SECTION: SESSION ANALYTICS
        st.markdown('<div class="section-header">🎯 SESSION ANALYTICS</div>', unsafe_allow_html=True)
        
        tab1, tab2, tab3 = st.tabs(["📊 FORMAT DISTRIBUTION", "🏫 DEPARTMENT ENGAGEMENT", "📅 UPCOMING SESSIONS"])
        
        with tab1:
            exploded_formats = filtered_df.explode('Session Format')['Session Format'].replace('', 'Not Specified')
            exploded_formats = exploded_formats[exploded_formats != 'Not Specified']
            if not exploded_formats.empty:
                format_counts = exploded_formats.value_counts()
                c1, c2 = st.columns(2)
                with c1:
                    fig_pie = px.pie(values=format_counts.values, names=format_counts.index, hole=0.4,
                                    color_discrete_sequence=TECHNO_COLORS['cyber'], title="Session Format Distribution")
                    fig_pie.update_traces(textposition='inside', textinfo='percent+label')
                    fig_pie.update_layout(height=400)
                    st.plotly_chart(fig_pie, use_container_width=True)
                with c2:
                    fig_bar = px.bar(x=format_counts.values, y=format_counts.index, orientation='h',
                                    color=format_counts.values, color_continuous_scale='Viridis',
                                    title="Session Format Frequency")
                    fig_bar.update_layout(height=400, xaxis_title="Count", yaxis_title="Format")
                    st.plotly_chart(fig_bar, use_container_width=True)
            else:
                st.info("No session format data available")
        
        with tab2:
            # Department Engagement
            st.markdown("### Alumni Count by Department")
            dept_counts = filtered_df['Department'].value_counts()
            if not dept_counts.empty:
                c1, c2 = st.columns(2)
                with c1:
                    fig_dept_pie = px.pie(values=dept_counts.values, names=dept_counts.index, hole=0.4,
                                        color_discrete_sequence=TECHNO_COLORS['neon'], 
                                        title="Department Distribution")
                    fig_dept_pie.update_traces(textposition='inside', textinfo='percent+label')
                    fig_dept_pie.update_layout(height=400)
                    st.plotly_chart(fig_dept_pie, use_container_width=True)
                with c2:
                    fig_dept_bar = px.bar(x=dept_counts.values, y=dept_counts.index, orientation='h',
                                        color=dept_counts.values, color_continuous_scale='Blues',
                                        title="Alumni Count by Department")
                    fig_dept_bar.update_layout(height=400, xaxis_title="Count", yaxis_title="Department")
                    st.plotly_chart(fig_dept_bar, use_container_width=True)
                    
                # Department engagement table
                dept_engagement = engagement_metrics['dept_engagement'].sort_values('Learning Hours', ascending=False)
                st.markdown("### Department Learning Engagement")
                st.dataframe(dept_engagement, use_container_width=True, height=200)
            else:
                st.info("No department data available")
        
        with tab3:
            if engagement_metrics['upcoming_sessions']:
                upcoming_df = pd.DataFrame(engagement_metrics['upcoming_sessions'])
                upcoming_df['date'] = pd.to_datetime(upcoming_df['date'])
                upcoming_df['formats'] = upcoming_df['formats'].apply(lambda x: ', '.join(x) if isinstance(x, list) else str(x))
                
                st.metric("Upcoming Sessions (Next 7 Days)", len(engagement_metrics['upcoming_sessions']))
                st.dataframe(upcoming_df[['date', 'alumni', 'department', 'formats', 'learning_hours']].rename(
                    columns={'date': 'Date', 'alumni': 'Alumni', 'department': 'Department', 
                            'formats': 'Formats', 'learning_hours': 'Learning Hours'}
                ), use_container_width=True, height=300)
            else:
                st.info("No upcoming sessions in the next 7 days")

        # SECTION: DETAILED DATA EXPLORER
        st.markdown('<div class="section-header">🔍 DETAILED DATA EXPLORER</div>', unsafe_allow_html=True)
        
        explorer_df = filtered_df.assign(**{
            'Session Formats': filtered_df['Session Format'].apply(lambda x: ', '.join(x) if isinstance(x, list) else str(x)),
            'Learning Hours': filtered_df['Learning Hours'].round(1)
        })[[
            'Name', 'Batch', 'Department', 'Position', 'Company', 'Format', 
            'Session Formats', 'Learning Hours', 'Session date'
        ]].sort_values('Learning Hours', ascending=False)
        
        st.dataframe(explorer_df, use_container_width=True, height=400)

        # SECTION: REPORT GENERATION
        st.markdown('<div class="section-header">📥 ADVANCED REPORT GENERATION</div>', unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # CSV Export
            csv = filtered_df.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="💾 DOWNLOAD CSV DATA",
                data=csv,
                file_name=f"alumni_analytics_{today}.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        with col2:
            try:
                word_buffer = generate_word_report(filtered_df, session_metrics, learning_metrics, engagement_metrics, today)
                st.download_button(
                    label="📝 DOWNLOAD WORD REPORT",
                    data=word_buffer,
                    file_name=f"alumni_analytics_report_{today}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Error generating Word document: {e}")

        # SECTION: QUICK STATS
        st.markdown('<div class="section-header">📈 QUICK STATS</div>', unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Avg Learning per Alumni", f"{learning_metrics['avg_learning_per_alumni']:.1f} hours")
        with col2:
            st.metric("Active Contributors", f"{learning_metrics['total_alumni_contributors']} alumni")
        with col3:
            st.metric("Remaining Sessions", f"{session_metrics['remaining_sessions']} sessions")
        with col4:
            st.metric("Unique Companies", f"{filtered_df['Company'].nunique()} companies")

    else:
        st.error("❌ FAILED TO LOAD DATA FROM GOOGLE SHEETS")

if __name__ == "__main__":
    main()
